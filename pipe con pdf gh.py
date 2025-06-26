import pandas as pd
import numpy as np
import unicodedata
import re
from rapidfuzz import fuzz
import matplotlib.pyplot as plt
import seaborn as sns
import os
from datetime import datetime
from docx import Document
import smtplib
from email.message import EmailMessage
from informe_pdf import generar_informe_pdf


# Cargar el archivo
df = pd.read_excel(r"C:\Users\primo\Downloads\Pedido de atención SM - CeSAC 24 (respuestas).xlsx", sheet_name="Respuestas de formulario 1")
# Setear ruta de salida y fecha de hoy
output_dir = r"C:\Users\primo\Desktop\Ciencia de Datos\Proyectos\Salud Mental\output"
fecha_hoy = datetime.now().strftime("%Y-%m-%d")
os.makedirs(output_dir, exist_ok=True)

# --- PROCESAMIENTO DE COLUMNA DE TEXTO LIBRE ---
# Diccionario de palabras clave
diccionario_palabras = {
    "angustia": [
        "angustia", "angustiado", "sensible", "sensible y angustiado"
    ],
    "depresión": [
        "depresión", "depresivo", "estado depresivo", "distimia", "ánimo triste", "desgano"
    ],
    "bullying": [
        "bullying", "acoso escolar"
    ],
    "conductas autolesivas": [
        "autolesión", "autolesiones", "conductas autolesivas", "autolesiva", "autolesivo", "autol"
    ],
    "ansiedad": [
        "ansiedad", "nervioso", "crisis de ansiedad", "crisis de ira", "ataques de ira", "temblores", "inquietud", "ansioso", "ansiosa"
    ],
    "obesidad": [
        "obesidad", "sobrepeso"
    ],
    "suicidio/ideas de muerte": [
        "suicidio", "suicida", "ideación suicida", "pensamientos suicidas", "pensamientos de muerte", "fantasías de muerte", "ideas de muerte", "ideas suicidas", "de muerte", "pensamientos de muerte", "fantasías de muerte", "pensamientos suicidas", "ideas de"
    ],
    "pánico": [
        "ataque de pánico", "crisis de pánico", "pánico"
    ],
    "violencia": [
        "violencia", "violento", "vínculo violento", "ex pareja violenta", "familiar violento","violenc","violent"
    ],
    "duelo": [
        "duelo", "pérdida", "perdió a", "falleció", "fallecieron",  "falleció", "fallecieron", "fallecimiento", "muerte", "de muerte", "muerte traumática", "duelo prolongado", "pérdida reciente", "duelo sin elaborar"
    ],
    "estrés": [
        "estrés", "estresado", "sobrecarga emocional", "sobrecarga", "agotamiento", "desbordada", "estresada", "sobrepasada"
    ],
    "discapacidad": [
        "discapacidad", "discapacitante", "silla de ruedas", "neuropatía", "trastorno cognoscitivo", "trastorno madurativo", "retraso madurativo", "retraso neuromadurativo", "retraso mental"
    ],
    "esquizofrenia": [
        "esquizofrenia", "psicosis", "trastornos psicóticos", "psicóticos", "descompensación psicótica"
    ],
    "consumo": [
        "consumo", "adicción", "adicciones", "alcohol", "alcohólico", "alcoholico", "problemas con el alcohol", "problemas con el juego", "juego patológico"
    ],
    "ASI": [
        "abuso", "abusos", "abuso sexual", "abuso físico", "violencia sexual", "asi"
    ],
    "TCA": [
        "trastorno alimentario", "trastorno de la conducta alimentaria", "conducta alimentaria", "tca", "anorexia", "bulimia", "desorden alimentario", "alimentación", "alimentaria", "alimenta"
    ],
    "enuresis": [
        "enuresis", "se hace pis"
    ],
    "encopresis": [
        "encopresis", "se hace caca"
    ],
    "problemas de sueño": [
        "trastorno del sueño", "dificultades para dormir", "alteración del sueño", "problemas para dormir", "insomnio", "problemas de sueño", "no puede dormir", "sueño"
    ],
    "problemas de conducta": [
        "problemas de conducta", "conducta desafiante", "desobediencia", "mal comportamiento", "mala conducta"
    ],
    "espectro autista": [
         "tea", "espectro autista", "asperger", "autismo", "autista"
    ],
    "continuidad de tratamiento": [
        "continuidad de tratamiento", "interrumpió tratamiento", "abandono del tratamiento", "cortaron prestación", "no pudo continuar", "perdida de seguimiento", "tto"
    ],
    "identidad de género": [
        "transición", "modificación hormonal", "tratamiento hormonal", "acompañamiento de género", "identidad de género"
    ],
    "tdah": [
        "tdah", "hiperactividad", "hiperactivo", "déficit de atención", "déficit atencional"
    ],
    "trayectoria escolar": [
        "abandono escolar", "dejo la escuela", "dejo el colegio", "no quiere ir a la escuela", "no quiere ir al colegio", "problemas escolares", "problemas en la escuela", "trayectoria escolar", "escuela", "colegio", "escolar"
    ]
}

# Función para normalizar texto
def normalizar(texto):
    if pd.isnull(texto):
        return ""
    texto = str(texto).lower()
    return unicodedata.normalize('NFKD', texto).encode('ASCII', 'ignore').decode('utf-8')

# Palabras clave normalizadas
diccionario_palabras_norm = {
    categoria: [normalizar(v) for v in variantes]
    for categoria, variantes in diccionario_palabras.items()
}
# Función de asignación de categorías con fuzzy match
def asignar_categorias(motivo, umbral=85):
    motivo_norm = normalizar(motivo)
    categorias = []

    for categoria, variantes in diccionario_palabras_norm.items():
        for variante in variantes:
            variante_norm = normalizar(variante)

            # Match exacto
            if re.search(rf'\b{re.escape(variante_norm)}\b', motivo_norm):
                categorias.append(categoria)
                break
            else:
                # Fuzzy match palabra por palabra
                for palabra_motivo in motivo_norm.split():
                    similitud = fuzz.ratio(variante_norm, palabra_motivo)
                    if similitud >= umbral:
                        categorias.append(categoria)
                        break
                else:
                    continue  # Solo se ejecuta si no hay break en inner loop
                break  # Rompe el loop de variantes si se encontró una

    return "/".join(sorted(set(categorias))) if categorias else None

# Aplicar al dataframe la función de asignación de categorías
df['palabras clave'] = df['Motivo del pedido'].apply(asignar_categorias)

# Filtrar los motivos sin palabras clave detectadas
sin_categoria = df[df['palabras clave'].isna()][['Motivo del pedido']]

# Exportarlos a un archivo Excel
archivo_sin_categoria = os.path.join(output_dir, f"motivos_sin_palabras_clave_{fecha_hoy}.xlsx")
sin_categoria.to_excel(archivo_sin_categoria, index=False)
print(f"Archivo exportado: {archivo_sin_categoria}")

# ---LIMPIEZA Y PREPARACIÓN DE DATOS---
# Limpieza de la columna 'Edad' para extraer el número
def extraer_edad(valor):
    if isinstance(valor, str):
        match = re.search(r'\d+', valor)
        if match:
            return int(match.group())
    elif isinstance(valor, (int, float)):
        return int(valor)
    return None

df['Edad_num'] = df['Edad'].apply(extraer_edad)

# Definir los bins y las etiquetas
bins = [0, 5, 11, 17, 49, float('inf')]
labels = ['0-5', '6-11', '12-17', '18-49', '50+']
# Crear la columna 'rango_edad' usando pd.cut
df['rango_edad'] = pd.cut(df['Edad_num'], bins=bins, labels=labels, right=True, include_lowest=True)

# Reemplazar valores específicos en la columna 'Lo que motiva la derivación es:'
df['Lo que motiva la derivación es:'] = df['Lo que motiva la derivación es:'].str.replace('Empeoramiento de un problema crónico. Violencia. Amenaza con tijeras.', 'Empeoramiento de un problema crónico', regex=False)
# Crear la columna 'mes' con el mes de la marca temporal
df['mes'] = df['Marca temporal'].dt.to_period('M')
# Crear una columna con la semana del año
df['semana'] = df['Marca temporal'].dt.to_period('W')
# Crear una función para quitar "angustia" si hay otras palabras clave
def quitar_angustia_si_hay_otras(palabras):
    if pd.isnull(palabras):
        return palabras
    claves = [k.strip() for k in palabras.split('/')]
    if 'angustia' in claves and len(claves) > 1:
        claves = [k for k in claves if k != 'angustia']
    return '/'.join(claves) if claves else None

df['palabras clave'] = df['palabras clave'].apply(quitar_angustia_si_hay_otras)
# Crear columna que cuenta la cantidad de palabras clave
df['cantidad_palabras_clave'] = df['palabras clave'].apply(lambda x: len(x.split('/')) if pd.notnull(x) else 0)

# --- VISUALIZACIÓN DE DATOS ---
# Configuración de estilo para las gráficas
sns.set_theme(palette="viridis")

# CREAR GRÁFICO DE BARRAS RANGO DE EDAD
# Obtener colores únicos para cada barra según cantidad de labels
colores = sns.color_palette("viridis", len(labels))

# Definir ruta de salida con fecha
archivo_rango_edad = os.path.join(output_dir, f"grafico_rango_edad_{fecha_hoy}.png")

plt.figure(figsize=(8, 5))
ax = sns.countplot(data=df, x='rango_edad', order=labels, palette=colores)

plt.xlabel('Rango de Edad')
plt.ylabel('Cantidad de Pedidos')
plt.title('Cantidad de pedidos por rango de edad')

# Etiquetas de datos encima de cada barra
for p in ax.patches:
    height = p.get_height()
    ax.annotate(f'{int(height)}', (p.get_x() + p.get_width() / 2, height),
                ha='center', va='bottom')

# Etiquetas eje x sin decimales
ax.set_xticklabels([str(label) for label in labels])
plt.tight_layout()

# Guardar gráfico y cerrar figura
plt.savefig(archivo_rango_edad, dpi=300)
plt.close()

print(f"Gráfico guardado: {archivo_rango_edad}")

# CREAR GRÁFICO CIRCULAR DE MOTIVOS DE DERIVACIÓN
# Conteo de motivos
motivos_counts = df['Lo que motiva la derivación es:'].value_counts()

# Obtener colores de la paleta viridis según cantidad de categorías
colores = sns.color_palette("viridis", len(motivos_counts))

plt.figure(figsize=(9, 9))
motivos_counts.plot.pie(
    autopct=lambda pct: f"{pct:.1f}%",
    startangle=90,
    counterclock=False,
    colors=colores
)
plt.ylabel('')
plt.title('Distribución de motivos de derivación')
plt.tight_layout()

# Guardar gráfico
archivo_pie = os.path.join(output_dir, f"grafico_circular_motivos_{fecha_hoy}.png")
plt.savefig(archivo_pie, dpi=300)
plt.close()

print(f"Gráfico circular guardado: {archivo_pie}")

# CREAR GRÁFICO CIRCULAR DE TRATAMIENTO PREVIO
# Conteo de respuestas sobre tratamiento previo
tratamiento_counts = df['¿La persona ya hizo tratamiento por salud mental alguna vez?'].value_counts()

# Colores viridis según cantidad de categorías
colores = sns.color_palette("viridis", len(tratamiento_counts))

plt.figure(figsize=(7, 7))
tratamiento_counts.plot.pie(
    autopct=lambda pct: f"{pct:.1f}%", 
    startangle=90, 
    counterclock=False,
    colors=colores
)

plt.ylabel('')
plt.title('¿La persona ya hizo tratamiento por salud mental alguna vez?')
plt.tight_layout()

# Guardar gráfico
archivo_tratamiento = os.path.join(output_dir, f"grafico_tratamiento_previo_{fecha_hoy}.png")
plt.savefig(archivo_tratamiento, dpi=300)
plt.close()

print(f"Gráfico circular guardado: {archivo_tratamiento}")

# CREAR GRÁFICO DE LÍNEA CANTIDAD DE PEDIDOS POR MES
# Agrupar por mes
df['mes'] = df['Marca temporal'].dt.to_period('M')
consultas_por_mes = df.groupby('mes').size()

# Convertir índice Period a datetime para control más preciso
consultas_por_mes.index = consultas_por_mes.index.to_timestamp()

# Tomar el color principal de viridis
color_linea = sns.color_palette("viridis", 1)[0]

plt.figure(figsize=(10, 5))
plt.plot(consultas_por_mes.index, consultas_por_mes.values, marker='o', color=color_linea)

# Extender ligeramente los límites del eje X
inicio = consultas_por_mes.index.min()
fin = consultas_por_mes.index.max()
delta = (fin - inicio) / 20  # Ajuste fino (~5%)
plt.xlim(inicio - delta, fin + delta)

plt.xlabel('Mes')
plt.ylabel('Cantidad de Pedidos')
plt.title('Cantidad de pedidos por mes')
plt.grid(True)
plt.tight_layout()

# Guardar gráfico
archivo_linea = os.path.join(output_dir, f"grafico_pedidos_por_mes_{fecha_hoy}.png")
plt.savefig(archivo_linea, dpi=300)
plt.close()

print(f"Gráfico de línea guardado: {archivo_linea}")

# CREAR GRÁFICO DE BARRAS APILADAS PEDIDOS POR MES Y SEMANA
# Crear tabla dinámica: filas=mes, columnas=semana, valores=cantidad de pedidos
pivot = df.pivot_table(index='mes', columns='semana', values='DNI', aggfunc='count', fill_value=0)

# Convertir índice y columnas a string para visualización
pivot.index = pivot.index.to_timestamp()
pivot.columns = pivot.columns.to_timestamp()

# Asignar número de semana dentro del mes a cada columna
semanas_ordenadas = sorted(pivot.columns)
semana_num = []

for semana in semanas_ordenadas:
    mes = semana.to_period('M')
    semanas_mes = sorted([s for s in semanas_ordenadas if s.to_period('M') == mes])
    semana_idx = semanas_mes.index(semana) + 1  # 1-based
    semana_num.append(semana_idx)

# Generar etiquetas "Semana 1", ..., "Semana 5"
max_semanas = max(semana_num)
semana_labels = [f"Semana {i}" for i in range(1, max_semanas + 1)]

# Obtener colores fijos para las semanas
colores = sns.color_palette("viridis", max_semanas)
colores_semanas = colores[:max_semanas]

# Crear gráfico
plt.figure(figsize=(12, 6))
bottom = np.zeros(len(pivot))

for i in range(max_semanas):
    # Filtrar columnas de la semana i+1
    cols_semana = [col for idx, col in enumerate(pivot.columns) if semana_num[idx] == i + 1]
    valores = pivot[cols_semana].sum(axis=1)
    plt.bar(pivot.index.astype(str), valores, bottom=bottom, color=colores_semanas[i], label=f"Semana {i + 1}")
    bottom += valores

plt.xlabel('Mes')
plt.ylabel('Cantidad de Pedidos')
plt.title('Pedidos por mes y semana')
plt.legend(title='Semana del mes')
plt.tight_layout()

# Guardar gráfico
archivo_apilado = os.path.join(output_dir, f"grafico_barras_pedidos_por_semana_{fecha_hoy}.png")
plt.savefig(archivo_apilado, dpi=300)
plt.close()

print(f"Gráfico de barras pedidos por semana guardado: {archivo_apilado}")

# GRÁFICO DE BARRAS DE CANTIDAD DE PALABRAS CLAVE POR REGISTRO
plt.figure(figsize=(10, 6))

# Ordenar valores únicos para el eje x
orden = sorted(df['cantidad_palabras_clave'].unique())

# Crear gráfico
ax = sns.countplot(
    data=df,
    x='cantidad_palabras_clave',
    order=orden,
    palette=sns.color_palette("viridis", len(orden))
)

# Títulos y ejes
plt.title('Cantidad de Palabras Clave por Registro')
plt.xlabel('Cantidad de Palabras Clave')
plt.ylabel('Cantidad de Registros')

# Etiquetas en cada barra
for p in ax.patches:
    height = p.get_height()
    ax.annotate(
        str(int(height)),
        (p.get_x() + p.get_width() / 2., height),
        ha='center', va='bottom',
        fontsize=10, color='black'
    )

# Guardar gráfico
archivo_barras = os.path.join(output_dir, f"cantidad_palabras_clave_{fecha_hoy}.png")
plt.tight_layout()
plt.savefig(archivo_barras, dpi=300)
plt.close()

print(f"Gráfico 'Cantidad de Palabras Clave por Registro' guardado: {archivo_barras}")

# CREAR GRÁFICO DE BARRAS DE TOP 10 MÓTIVOS MÁS FRECUENTES
grupo = df.groupby('palabras clave')['DNI'].count()
grupo = grupo.sort_values(ascending=False)
categorías = grupo.rename_axis('categoría').reset_index(name='cantidad')
# Paso 1: separar las categorías múltiples en filas individuales
categorías_expandidas = (
    categorías
    .assign(categoría=categorías['categoría'].str.split('/'))  # convierte a lista
    .explode('categoría')                                      # cada categoría individual como fila
)

# Paso 2: agrupar y sumar las cantidades por categoría individual
categorías_agrupadas = (
    categorías_expandidas
    .groupby('categoría', as_index=False)['cantidad']
    .sum()
    .sort_values('cantidad', ascending=False)
)
categorías_agrupadas['categoría'] = categorías_agrupadas['categoría'].str.capitalize()
top_10 = categorías_agrupadas.sort_values('cantidad', ascending=False).head(10)
# Configurar el gráfico
plt.figure(figsize=(10, 6))

# Paleta viridis para el top 10
colores_top10 = sns.color_palette("viridis", len(top_10))

# Crear gráfico
ax = sns.barplot(
    data=top_10,
    x='cantidad',
    y='categoría',
    palette=colores_top10
)

# Títulos y etiquetas
plt.title('Top 10 motivos más frecuentes', fontsize=14)
plt.xlabel('Cantidad de casos')
plt.ylabel('Motivo')

# Etiquetas en las barras
for i, v in enumerate(top_10['cantidad']):
    ax.text(v + 0.5, i, str(v), va='center', fontsize=10)

# Guardar gráfico
archivo_top10 = os.path.join(output_dir, f"top10_motivos_{fecha_hoy}.png")
plt.tight_layout()
plt.savefig(archivo_top10, dpi=300)
plt.close()

print(f"Gráfico 'Top 10 motivos más frecuentes' guardado: {archivo_top10}")
# Exportar las categorías procesadas a un archivo Excel
archivo_categorias = os.path.join(output_dir, f"categorias_agrupadas_{fecha_hoy}.xlsx")
categorías_agrupadas.to_excel(archivo_categorias, index=False)

print(f"Archivo 'categorias_agrupadas_{fecha_hoy}.xlsx' exportado correctamente en: {archivo_categorias}")

# --- CREAR ARCHIVOS RESUMEN ---
# Preparar contenido del resumen
semana = df.groupby('semana').count()['DNI'].reset_index()
semana = semana.rename(columns={'DNI': 'Cantidad de consultas'})
promedio_consultas_semanales = semana['Cantidad de consultas'].mean()
# Crear documento Word
doc = Document()
doc.add_heading(f"Resumen de métricas del análisis de salud mental - {fecha_hoy}", level=1)

# Agregar métricas principales
doc.add_paragraph(f"Promedio de consultas semanales: {promedio_consultas_semanales:.2f}")
doc.add_paragraph(f"Cantidad total de registros analizados: {len(df)}")
doc.add_paragraph(f"Cantidad de registros sin palabras clave: {len(sin_categoria)}")
doc.add_paragraph(f"Cantidad promedio de palabras clave por registro: {df['cantidad_palabras_clave'].mean():.2f}")
doc.add_paragraph(f"Cantidad total de categorías definidas: {len(diccionario_palabras)}")

# Título para el listado de palabras clave
doc.add_heading("Listado de categorías y variantes asociadas:", level=2)

# Agregar cada categoría como un párrafo separado
for clave, variantes in diccionario_palabras.items():
    doc.add_paragraph(f"{clave}: {', '.join(variantes)}")

# Guardar el archivo
archivo_doc = os.path.join(output_dir, f"resumen_metricas_{fecha_hoy}.docx")
doc.save(archivo_doc)

print(f"Archivo de resumen en Word creado: {archivo_doc}")

# Crear informe PDF
generar_informe_pdf(df, diccionario_palabras, output_dir, fecha_hoy)
print(f"Informe PDF creado en: {output_dir}")

# --- ENVIAR ANÁLISIS COMPLETO ---
# Configuración
GMAIL_USER = 'leonel.toro93@gmail.com'
GMAIL_APP_PASSWORD = '******'  # Reemplaza con tu contraseña de aplicación de Gmail
to_email = 'destinatario@.com'  # Reemplaza con el correo del destinatario

# Crear el mensaje
msg = EmailMessage()
msg['Subject'] = 'Análisis Salud Mental - Envío automático'
msg['From'] = GMAIL_USER
msg['To'] = to_email
msg.set_content("Envío análisis del formulario de pedidos de atención de salud mental")

# Adjuntar archivos de output_dir
for filename in os.listdir(output_dir):
    filepath = os.path.join(output_dir, filename)
    if os.path.isfile(filepath):
        with open(filepath, 'rb') as f:
            file_data = f.read()
            file_name = os.path.basename(filepath)
        # Intentar detectar el tipo MIME según extensión
        if filename.lower().endswith('.xlsx'):
            maintype, subtype = 'application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif filename.lower().endswith('.png'):
            maintype, subtype = 'image', 'png'
        elif filename.lower().endswith('.jpg') or filename.lower().endswith('.jpeg'):
            maintype, subtype = 'image', 'jpeg'
        elif filename.lower().endswith('.docx'):
            maintype, subtype = 'application', 'vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.lower().endswith('.txt'):
            maintype, subtype = 'text', 'plain'
        else:
            maintype, subtype = 'application', 'octet-stream'  # genérico

        msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=file_name)

# Enviar el correo
with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
    smtp.login(GMAIL_USER, GMAIL_APP_PASSWORD)
    smtp.send_message(msg)

print("Correo enviado con éxito.")