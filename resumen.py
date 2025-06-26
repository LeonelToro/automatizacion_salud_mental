# resumen.py
import os
from docx import Document

def crear_resumen_metricas(df, sin_categoria, diccionario_palabras, output_dir, fecha_hoy):
    semana = df.groupby('semana').count()['DNI'].reset_index()
    semana = semana.rename(columns={'DNI': 'Cantidad de consultas'})
    promedio_consultas_semanales = semana['Cantidad de consultas'].mean()

    doc = Document()
    doc.add_heading(f"Resumen de métricas del análisis de salud mental - {fecha_hoy}", level=1)

    doc.add_paragraph(f"Promedio de consultas semanales: {promedio_consultas_semanales:.2f}")
    doc.add_paragraph(f"Cantidad total de registros analizados: {len(df)}")
    doc.add_paragraph(f"Cantidad de registros sin palabras clave: {len(sin_categoria)}")
    doc.add_paragraph(f"Cantidad promedio de palabras clave por registro: {df['cantidad_palabras_clave'].mean():.2f}")
    doc.add_paragraph(f"Cantidad total de categorías definidas: {len(diccionario_palabras)}")

    doc.add_heading("Listado de categorías y variantes asociadas:", level=2)

    for clave, variantes in diccionario_palabras.items():
        doc.add_paragraph(f"{clave}: {', '.join(variantes)}")

    archivo_doc = os.path.join(output_dir, f"resumen_metricas_{fecha_hoy}.docx")
    doc.save(archivo_doc)
    print(f"Resumen guardado: {archivo_doc}")
